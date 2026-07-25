"""Enhance the educational Smart Shopping AI Agent Word report.

The script preserves the source document and writes a new DOCX containing:
- a visual system architecture diagram;
- concise design-decision and structured-handoff explanations;
- an educational evaluation matrix based on the offline test suite; and
- corrected Markdown/PDF output references and section numbering.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import sys
import textwrap

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


NAVY = "17375E"
MEDIUM_BLUE = "315B7D"
LIGHT_BLUE = "DCE8F2"
PALE_BLUE = "EEF4F8"
WHITE = "FFFFFF"
TEXT = "243447"


def find_paragraph(document: Document, exact_text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Could not find paragraph: {exact_text!r}")


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace text while retaining the paragraph's style."""

    paragraph.clear()
    paragraph.add_run(new_text)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        margin = margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def format_table(table, widths: tuple[float, ...] | None = None) -> None:
    """Apply the report's navy and pale-blue table appearance."""

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        if row_index == 0:
            repeat_table_header(row)
        fill = NAVY if row_index == 0 else (LIGHT_BLUE if row_index % 2 else WHITE)
        for column_index, cell in enumerate(row.cells):
            set_cell_shading(cell, fill)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and column_index < len(widths):
                cell.width = Inches(widths[column_index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_table(document: Document, rows: list[list[str]], widths: tuple[float, ...]):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            table.cell(row_index, column_index).text = value
    format_table(table, widths)
    return table


def insert_after(anchor, block):
    """Move a paragraph or table directly after an XML block element."""

    element = block._p if hasattr(block, "_p") else block._tbl
    anchor.addnext(element)
    return element


def create_paragraph(document: Document, text: str = "", style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    return paragraph


def add_callout(document: Document, text: str):
    table = document.add_table(rows=1, cols=1)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(TEXT)
    return table


def load_font(size: int, bold: bool = False):
    filename = "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"
    path = Path("/usr/share/fonts/truetype/dejavu") / filename
    return ImageFont.truetype(str(path), size=size)


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    subtitle: str,
    title_font,
    subtitle_font,
    title_fill: str,
    subtitle_fill: str,
) -> None:
    x1, y1, x2, y2 = box
    width = x2 - x1
    title_lines = wrapped_lines(draw, title, title_font, width - 70)
    subtitle_lines = wrapped_lines(draw, subtitle, subtitle_font, width - 70)
    title_metrics = [
        (line, draw.textbbox((0, 0), line, font=title_font))
        for line in title_lines
    ]
    subtitle_metrics = [
        (line, draw.textbbox((0, 0), line, font=subtitle_font))
        for line in subtitle_lines
    ]
    title_height = sum(
        bounds[3] - bounds[1] + 5 for _, bounds in title_metrics
    )
    subtitle_height = sum(
        bounds[3] - bounds[1] + 3 for _, bounds in subtitle_metrics
    )
    gap = 12 if subtitle_lines else 0
    y = y1 + ((y2 - y1) - title_height - subtitle_height - gap) / 2

    for line, bounds in title_metrics:
        draw.text(
            (
                x1 + (width - (bounds[2] - bounds[0])) / 2,
                y - bounds[1],
            ),
            line,
            font=title_font,
            fill=title_fill,
        )
        y += bounds[3] - bounds[1] + 5
    y += gap
    for line, bounds in subtitle_metrics:
        draw.text(
            (
                x1 + (width - (bounds[2] - bounds[0])) / 2,
                y - bounds[1],
            ),
            line,
            font=subtitle_font,
            fill=subtitle_fill,
        )
        y += bounds[3] - bounds[1] + 3


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    subtitle: str,
    *,
    fill: str,
    outline: str,
    title_color: str,
    subtitle_color: str,
    title_font,
    subtitle_font,
    radius: int = 24,
    width: int = 5,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    draw_centered_text(
        draw,
        box,
        title,
        subtitle,
        title_font,
        subtitle_font,
        title_color,
        subtitle_color,
    )


def draw_down_arrow(draw: ImageDraw.ImageDraw, center_x: int, top: int, bottom: int) -> None:
    arrow_tip = bottom
    line_bottom = bottom - 24
    draw.line((center_x, top, center_x, line_bottom), fill=f"#{MEDIUM_BLUE}", width=9)
    draw.polygon(
        [
            (center_x - 20, line_bottom),
            (center_x + 20, line_bottom),
            (center_x, arrow_tip),
        ],
        fill=f"#{MEDIUM_BLUE}",
    )


def create_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1800), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(39, bold=True)
    body_font = load_font(29)
    small_title_font = load_font(27, bold=True)
    small_body_font = load_font(23)
    pill_font = load_font(27, bold=True)

    main_left, main_right = 420, 1380
    center_x = 900

    draw_box(
        draw,
        (main_left, 55, main_right, 160),
        "User Shopping Request",
        "{product_query}: product, budget, preferences, and location",
        fill=f"#{NAVY}",
        outline=f"#{NAVY}",
        title_color="white",
        subtitle_color="white",
        title_font=title_font,
        subtitle_font=body_font,
    )
    draw_down_arrow(draw, center_x, 160, 225)

    draw_box(
        draw,
        (main_left, 225, main_right, 355),
        "Runtime Checks and Guardrails",
        "Validate request scope and required API keys before paid work",
        fill=f"#{PALE_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )
    draw_down_arrow(draw, center_x, 355, 425)

    draw_box(
        draw,
        (main_left, 425, main_right, 570),
        "1. Product Search Specialist",
        "Finds candidates and verifies selected product pages",
        fill=f"#{NAVY}",
        outline=f"#{NAVY}",
        title_color="white",
        subtitle_color="white",
        title_font=title_font,
        subtitle_font=body_font,
    )
    draw_box(
        draw,
        (25, 430, 350, 565),
        "Research Tools",
        "Exa Search + website scraper",
        fill=f"#{LIGHT_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )
    draw.line((350, 497, main_left, 497), fill=f"#{MEDIUM_BLUE}", width=7)
    draw.polygon([(main_left - 20, 480), (main_left - 20, 514), (main_left, 497)], fill=f"#{MEDIUM_BLUE}")
    draw_down_arrow(draw, center_x, 570, 625)

    draw_box(
        draw,
        (610, 625, 1190, 700),
        "ProductSearchResult",
        "Validated evidence handoff",
        fill=f"#{LIGHT_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=pill_font,
        subtitle_font=small_body_font,
        radius=38,
    )
    draw_down_arrow(draw, center_x, 700, 765)

    draw_box(
        draw,
        (main_left, 765, main_right, 910),
        "2. Price Comparison Analyst",
        "Compares known costs, value, warranties, and limitations",
        fill=f"#{NAVY}",
        outline=f"#{NAVY}",
        title_color="white",
        subtitle_color="white",
        title_font=title_font,
        subtitle_font=body_font,
    )
    draw_box(
        draw,
        (25, 770, 350, 905),
        "Optional Tool",
        "One targeted Exa verification search",
        fill=f"#{LIGHT_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )
    draw.line((350, 837, main_left, 837), fill=f"#{MEDIUM_BLUE}", width=7)
    draw.polygon([(main_left - 20, 820), (main_left - 20, 854), (main_left, 837)], fill=f"#{MEDIUM_BLUE}")
    draw_down_arrow(draw, center_x, 910, 965)

    draw_box(
        draw,
        (610, 965, 1190, 1040),
        "PriceComparisonResult",
        "Validated comparison handoff",
        fill=f"#{LIGHT_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=pill_font,
        subtitle_font=small_body_font,
        radius=38,
    )
    draw_down_arrow(draw, center_x, 1040, 1105)

    draw_box(
        draw,
        (main_left, 1105, main_right, 1250),
        "3. Product Recommendation Advisor",
        "Selects the best-supported option and alternatives",
        fill=f"#{NAVY}",
        outline=f"#{NAVY}",
        title_color="white",
        subtitle_color="white",
        title_font=title_font,
        subtitle_font=body_font,
    )
    draw_box(
        draw,
        (1450, 1110, 1775, 1245),
        "No New Search",
        "Uses prior structured context only",
        fill=f"#{LIGHT_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )
    draw.line((main_right, 1177, 1450, 1177), fill=f"#{MEDIUM_BLUE}", width=7)
    draw.polygon([(main_right + 20, 1160), (main_right + 20, 1194), (main_right, 1177)], fill=f"#{MEDIUM_BLUE}")
    draw_down_arrow(draw, center_x, 1250, 1305)

    draw_box(
        draw,
        (580, 1305, 1220, 1385),
        "FinalRecommendationResult",
        "Validated JSON recommendation",
        fill=f"#{LIGHT_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=pill_font,
        subtitle_font=small_body_font,
        radius=38,
    )
    draw_down_arrow(draw, center_x, 1385, 1450)

    draw_box(
        draw,
        (main_left, 1450, main_right, 1565),
        "Deterministic Report Renderer",
        "One validated result creates both report formats",
        fill=f"#{PALE_BLUE}",
        outline=f"#{MEDIUM_BLUE}",
        title_color=f"#{NAVY}",
        subtitle_color=f"#{TEXT}",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )

    draw.line((center_x, 1565, center_x, 1605), fill=f"#{MEDIUM_BLUE}", width=9)
    draw.line((620, 1605, 1180, 1605), fill=f"#{MEDIUM_BLUE}", width=9)
    draw_down_arrow(draw, 620, 1605, 1640)
    draw_down_arrow(draw, 1180, 1605, 1640)
    draw_box(
        draw,
        (390, 1640, 850, 1755),
        "Markdown Report",
        "outputs/report.md",
        fill=f"#{NAVY}",
        outline=f"#{NAVY}",
        title_color="white",
        subtitle_color="white",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )
    draw_box(
        draw,
        (950, 1640, 1410, 1755),
        "PDF Report",
        "outputs/report.pdf",
        fill=f"#{NAVY}",
        outline=f"#{NAVY}",
        title_color="white",
        subtitle_color="white",
        title_font=small_title_font,
        subtitle_font=small_body_font,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", optimize=True)


def update_existing_content(document: Document) -> None:
    executive = find_paragraph(
        document,
        "The system is built with CrewAI and uses three specialized agents. "
        "The first agent searches for products that match the user’s request "
        "and verifies selected pages. The second agent compares the available "
        "prices, specifications, sellers, warranties, delivery details, and "
        "overall value. The final agent selects the best-supported option. "
        "The tasks exchange validated Pydantic objects, and a deterministic "
        "Python callback renders the final result as a Markdown recommendation "
        "with alternatives and source URLs.",
    )
    replace_paragraph_text(
        executive,
        "The system is built with CrewAI and uses three specialized agents. "
        "The first agent searches for products that match the user’s request "
        "and verifies selected pages. The second agent compares the available "
        "prices, specifications, sellers, warranties, delivery details, and "
        "overall value. The final agent selects the best-supported option. "
        "The tasks exchange validated Pydantic objects, and a deterministic "
        "Python callback renders the final result as matching Markdown and PDF "
        "reports with alternatives and source URLs.",
    )

    output_paragraph = find_paragraph(
        document,
        "The final advisor produces a validated FinalRecommendationResult. "
        "A deterministic Python callback then renders that structured data "
        "into the Markdown file `report.md`. The format is readable by a "
        "student, teacher, or buyer without requiring knowledge of the "
        "underlying task schemas.",
    )
    replace_paragraph_text(
        output_paragraph,
        "The final advisor produces a validated FinalRecommendationResult. "
        "A deterministic Python callback then renders the same structured data "
        "into `outputs/report.md` and `outputs/report.pdf`. Both formats contain "
        "the same recommendation and sources, and PDF generation does not require "
        "another agent or model call.",
    )

    heading_updates = {
        "11. Safe-Buying Guidance": "12. Safe-Buying Guidance",
        "12. Current Limitations": "13. Current Limitations",
        "13. Future Improvements": "14. Future Improvements",
        "14. Conclusion": "15. Conclusion",
    }
    for old_text, new_text in heading_updates.items():
        replace_paragraph_text(find_paragraph(document, old_text), new_text)
    find_paragraph(document, "15. Conclusion").paragraph_format.page_break_before = True

    for table in document.tables:
        for row in table.rows:
            key = row.cells[0].text.strip()
            if key == "Main Output":
                row.cells[1].text = "Evidence-based Markdown and PDF purchasing reports"
            elif key == "11–14":
                row.cells[0].text = "11–15"
                row.cells[1].text = (
                    "Evaluation, safe buying, limitations, improvements, and conclusion"
                )
            elif key == "05":
                row.cells[1].text = (
                    "A Python callback renders the structured result to "
                    "outputs/report.md and outputs/report.pdf."
                )
            elif key == "Output file":
                row.cells[0].text = "Output files"
                row.cells[1].text = "outputs/report.md; outputs/report.pdf"
            elif key == "Report renderer":
                row.cells[1].text = "tools.reporting.save_recommendation_report"

    technology_table = document.tables[5]
    if not any(row.cells[0].text.strip() == "PyMuPDF" for row in technology_table.rows):
        cells = technology_table.add_row().cells
        cells[0].text = "PyMuPDF"
        cells[1].text = "Renders the validated recommendation as a printable PDF."
        cells[2].text = "Creates a shareable report without another model call."
        row_index = len(technology_table.rows) - 1
        fill = LIGHT_BLUE if row_index % 2 else WHITE
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def insert_architecture_sections(document: Document, diagram_path: Path) -> None:
    workflow_table = next(
        table
        for table in document.tables
        if table.cell(0, 0).text.strip() == "01"
        and "shopping request" in table.cell(0, 1).text
    )
    anchor = workflow_table._tbl

    heading = create_paragraph(document, "6.1 Architecture Overview", "Heading 2")
    anchor = insert_after(anchor, heading)

    explanation = create_paragraph(
        document,
        "The diagram below shows both the flow of information and the separation "
        "of responsibilities. Search tools are available only where new evidence "
        "is needed, while each completed stage passes a typed result to the next "
        "agent.",
        "Normal",
    )
    anchor = insert_after(anchor, explanation)

    image_paragraph = create_paragraph(document)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_together = True
    run = image_paragraph.add_run()
    run.add_picture(str(diagram_path), width=Inches(6.25))
    doc_properties = run._r.xpath(".//wp:docPr")
    if doc_properties:
        doc_properties[0].set(
            "descr",
            "Architecture diagram showing the user request, runtime checks, "
            "three sequential CrewAI agents, Pydantic handoffs, and Markdown "
            "and PDF report outputs.",
        )
        doc_properties[0].set("title", "Smart Shopping AI Agent Architecture")
    anchor = insert_after(anchor, image_paragraph)

    caption = create_paragraph(
        document,
        "Figure 1. System architecture and structured data flow.",
        "Caption",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor = insert_after(anchor, caption)

    heading = create_paragraph(document, "6.2 Key Design Decisions", "Heading 2")
    anchor = insert_after(anchor, heading)

    decisions = [
        (
            "Specialized roles: each agent has one clear responsibility, making "
            "the workflow easier to understand, test, and explain."
        ),
        (
            "Sequential execution: comparison depends on research evidence, and "
            "recommendation depends on both earlier results."
        ),
        (
            "Structured handoffs: Pydantic contracts reduce missing fields and "
            "prevent free-form reports from becoming the next agent’s only input."
        ),
        (
            "Restricted tool access: the advisor cannot start a conflicting search, "
            "and the analyst uses search only for targeted verification."
        ),
        (
            "Deterministic reporting: Python creates Markdown and PDF from one "
            "validated result, avoiding an additional model call."
        ),
    ]
    for decision in decisions:
        paragraph = create_paragraph(document, decision, "List Bullet")
        anchor = insert_after(anchor, paragraph)

    heading = create_paragraph(document, "6.3 Structured Agent Handoffs", "Heading 2")
    anchor = insert_after(anchor, heading)

    paragraph = create_paragraph(
        document,
        "The contracts make the educational purpose of each stage visible. They "
        "validate field names and data types while keeping uncertain facts explicit.",
        "Normal",
    )
    anchor = insert_after(anchor, paragraph)

    handoff_table = add_table(
        document,
        [
            ["Stage", "Structured Contract", "Information Passed Forward"],
            [
                "Research",
                "ProductSearchResult",
                "Requirements, product evidence, source URLs, research issues, and shortfall",
            ],
            [
                "Comparison",
                "PriceComparisonResult",
                "Per-product costs, value analysis, price ranges, best-value deal, and limitations",
            ],
            [
                "Recommendation",
                "FinalRecommendationResult",
                "Selected option, alternatives, confidence, checklist, and numbered sources",
            ],
        ],
        widths=(1.05, 1.75, 3.35),
    )
    anchor = insert_after(anchor, handoff_table)

    callout = add_callout(
        document,
        "Simplified handoff example\n"
        "ProductSearchResult(products=[…])  →  "
        "PriceComparisonResult(comparisons=[…])  →  "
        "FinalRecommendationResult(recommendation=…, confidence=…)",
    )
    insert_after(anchor, callout)


def insert_evaluation_section(document: Document) -> None:
    evaluation_label = find_paragraph(document, "EVALUATION")
    anchor = evaluation_label._p

    heading = create_paragraph(document, "11. Educational Evaluation", "Heading 1")
    anchor = insert_after(anchor, heading)

    paragraph = create_paragraph(
        document,
        "The offline test suite was run on 25 July 2026 using "
        "`python -m unittest discover -s tests -v`. All 27 tests passed without "
        "using Exa Search or DeepSeek credits. The tests demonstrate that the "
        "project is configured correctly and handles its structured outputs "
        "consistently.",
        "Normal",
    )
    anchor = insert_after(anchor, paragraph)

    evaluation_table = add_table(
        document,
        [
            ["Evaluation Area", "Evidence Checked", "Result"],
            [
                "Crew configuration",
                "Three agents, tools, models, task order, and context references",
                "Passed",
            ],
            [
                "Runtime guardrails",
                "Missing keys, prohibited requests, vague requests, and valid scoped requests",
                "Passed",
            ],
            [
                "Structured data",
                "Exact decimals, HTTP(S) URLs, extra-field rejection, and five-product limit",
                "Passed",
            ],
            [
                "Crew loading",
                "Schemas and final report callback resolve from the JSONC configuration",
                "Passed",
            ],
            [
                "Report generation",
                "Readable Markdown and PDF, source links, and invalid JSON rejection",
                "Passed",
            ],
            [
                "Live factual accuracy",
                "Prices, stock, seller claims, and delivery information can change online",
                "Manual verification required",
            ],
        ],
        widths=(1.35, 3.65, 1.25),
    )
    anchor = insert_after(anchor, evaluation_table)

    note = add_callout(
        document,
        "Evaluation interpretation\n"
        "Passing tests shows that the software structure, validation, and report "
        "generation work as intended. It does not prove that every live shopping "
        "claim is permanently correct, so source links and the final buying "
        "checklist remain important.",
    )
    insert_after(anchor, note)


def main() -> None:
    if len(sys.argv) not in {1, 3}:
        raise SystemExit(
            "Usage: update_educational_report.py [SOURCE.docx OUTPUT.docx]"
        )

    default_source = Path(
        "/home/reaksmey-rin/Downloads/"
        "Smart_Shopping_Price_Comparison_AI_Agent_Report_Reaksmey_Rin_Updated.docx"
    )
    default_output = Path(
        "deliverables/"
        "Smart_Shopping_Price_Comparison_AI_Agent_Report_"
        "Reaksmey_Rin_Educational_Enhanced.docx"
    )
    source = Path(sys.argv[1]) if len(sys.argv) == 3 else default_source
    output = Path(sys.argv[2]) if len(sys.argv) == 3 else default_output
    diagram_path = output.parent / "smart_shopping_ai_agent_architecture.png"

    output.parent.mkdir(parents=True, exist_ok=True)
    create_architecture_diagram(diagram_path)

    document = Document(source)
    update_existing_content(document)
    insert_architecture_sections(document, diagram_path)
    insert_evaluation_section(document)

    document.core_properties.title = (
        "Smart Shopping & Price-Comparison AI Agent — Educational Project Report"
    )
    document.core_properties.subject = (
        "CrewAI multi-agent architecture, implementation, and educational evaluation"
    )
    document.core_properties.modified = datetime.now(timezone.utc)
    document.save(output)
    print(output.resolve())
    print(diagram_path.resolve())


if __name__ == "__main__":
    main()
